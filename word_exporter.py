from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from typing import Dict, List
import os
from datetime import datetime
from config import Config

class WordExporter:
    def __init__(self):
        self.document = Document()
        self.setup_document_styles()
    
    def setup_document_styles(self):
        """Setup professional screenplay formatting styles"""
        # Set page margins for screenplay format
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1)
        
        # Create custom styles
        styles = self.document.styles
        
        # Scene heading style
        scene_style = styles.add_style('Scene Heading', WD_STYLE_TYPE.PARAGRAPH)
        scene_style.font.name = 'Courier New'
        scene_style.font.size = Pt(12)
        scene_style.font.bold = True
        scene_style.font.all_caps = True
        
        # Action style
        action_style = styles.add_style('Action', WD_STYLE_TYPE.PARAGRAPH)
        action_style.font.name = 'Courier New'
        action_style.font.size = Pt(12)
        action_style.paragraph_format.left_indent = Inches(0)
        action_style.paragraph_format.right_indent = Inches(0)
        
        # Character style
        character_style = styles.add_style('Character', WD_STYLE_TYPE.PARAGRAPH)
        character_style.font.name = 'Courier New'
        character_style.font.size = Pt(12)
        character_style.font.bold = True
        character_style.paragraph_format.left_indent = Inches(3.5)
        character_style.paragraph_format.right_indent = Inches(2)
        
        # Dialogue style
        dialogue_style = styles.add_style('Dialogue', WD_STYLE_TYPE.PARAGRAPH)
        dialogue_style.font.name = 'Courier New'
        dialogue_style.font.size = Pt(12)
        dialogue_style.paragraph_format.left_indent = Inches(2.5)
        dialogue_style.paragraph_format.right_indent = Inches(2.5)
        
        # Parenthetical style
        parenthetical_style = styles.add_style('Parenthetical', WD_STYLE_TYPE.PARAGRAPH)
        parenthetical_style.font.name = 'Courier New'
        parenthetical_style.font.size = Pt(12)
        parenthetical_style.font.italic = True
        parenthetical_style.paragraph_format.left_indent = Inches(3)
        parenthetical_style.paragraph_format.right_indent = Inches(2.5)
    
    def add_title_page(self, title: str, author: str = "Screenwriter"):
        """Add a professional title page"""
        # Title
        title_para = self.document.add_paragraph()
        title_run = title_para.add_run(title.upper())
        title_run.font.name = 'Courier New'
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some spacing
        self.document.add_paragraph()
        self.document.add_paragraph()
        
        # Author
        author_para = self.document.add_paragraph()
        author_run = author_para.add_run(f"by\n{author}")
        author_run.font.name = 'Courier New'
        author_run.font.size = Pt(14)
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page break
        self.document.add_page_break()
    
    def add_scene(self, scene_data: Dict):
        """Add a scene in professional screenplay format"""
        # Scene heading
        scene_number = scene_data.get('scene_number', 'N/A')
        setting = scene_data.get('setting', 'Unknown Location')
        scene_heading = f"SCENE {scene_number} - {setting.upper()}"
        
        scene_para = self.document.add_paragraph(scene_heading, style='Scene Heading')
        
        # Action/Description
        action = scene_data.get('action', '')
        if action:
            action_para = self.document.add_paragraph(action, style='Action')
        
        # Characters and Dialogue
        characters = scene_data.get('characters', '')
        dialogue = scene_data.get('dialogue', '')
        
        if characters and dialogue:
            # Split characters if multiple
            char_list = [char.strip() for char in characters.split(',')]
            
            # Split dialogue into character parts (simple approach)
            dialogue_parts = dialogue.split('\n')
            
            for i, char in enumerate(char_list):
                if i < len(dialogue_parts):
                    # Character name
                    char_para = self.document.add_paragraph(char.upper(), style='Character')
                    
                    # Dialogue
                    dialogue_text = dialogue_parts[i].strip()
                    if dialogue_text:
                        dialogue_para = self.document.add_paragraph(dialogue_text, style='Dialogue')
        
        # Notes (if any)
        notes = scene_data.get('notes', '')
        if notes:
            notes_para = self.document.add_paragraph()
            notes_run = notes_para.add_run(f"NOTES: {notes}")
            notes_run.font.name = 'Courier New'
            notes_run.font.size = Pt(10)
            notes_run.font.italic = True
            notes_para.paragraph_format.left_indent = Inches(0.5)
        
        # Add spacing between scenes
        self.document.add_paragraph()
    
    def add_character_list(self, characters: List[Dict]):
        """Add a character list at the end"""
        self.document.add_page_break()
        
        # Character List heading
        char_heading = self.document.add_paragraph("CHARACTER LIST", style='Scene Heading')
        
        for char in characters:
            char_name = char.get('name', 'Unknown')
            char_age = char.get('age', 'Unknown')
            char_desc = char.get('description', 'No description')
            char_personality = char.get('personality', 'No personality traits')
            
            char_para = self.document.add_paragraph()
            char_run = char_para.add_run(f"{char_name.upper()} (Age: {char_age})")
            char_run.font.name = 'Courier New'
            char_run.font.size = Pt(12)
            char_run.font.bold = True
            
            desc_para = self.document.add_paragraph(char_desc)
            desc_para.style = self.document.styles['Action']
            
            personality_para = self.document.add_paragraph(char_personality)
            personality_para.style = self.document.styles['Action']
            
            self.document.add_paragraph()
    
    def export_scenes_to_word(self, scenes: List[Dict], characters: List[Dict], 
                            title: str = "Screenplay", author: str = "Screenwriter") -> str:
        """Export scenes to a Word document"""
        try:
            # Add title page
            self.add_title_page(title, author)
            
            # Add scenes
            for scene in scenes:
                self.add_scene(scene)
            
            # Add character list
            if characters:
                self.add_character_list(characters)
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{title.replace(' ', '_')}_{timestamp}.docx"
            filepath = os.path.join(Config.SCRIPT_FILE_PATH, filename)
            
            # Save document
            self.document.save(filepath)
            
            return filepath
            
        except Exception as e:
            raise Exception(f"Error exporting to Word: {str(e)}")
    
    def export_single_scene(self, scene_data: Dict, title: str = "Scene") -> str:
        """Export a single scene to Word"""
        try:
            # Add title page
            self.add_title_page(f"{title} - Scene {scene_data.get('scene_number', 'N/A')}")
            
            # Add the scene
            self.add_scene(scene_data)
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Scene_{scene_data.get('scene_number', 'N/A')}_{timestamp}.docx"
            filepath = os.path.join(Config.SCRIPT_FILE_PATH, filename)
            
            # Save document
            self.document.save(filepath)
            
            return filepath
            
        except Exception as e:
            raise Exception(f"Error exporting scene to Word: {str(e)}") 